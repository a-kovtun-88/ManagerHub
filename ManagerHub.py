from docx import Document
from docx.shared import Inches

commercial_number = input("Введите номер коммерческого предложения >> ")
commercial_date = input("Введите дату коммерческого предложения >> ")
commercial_name = input("Введите название коммерческого предложения >> ")

#БАЗОВЫЕ ПАРАМЕТРЫ

bk_count = float(input("Введите количество блок-контейнеров, шт. >> "))
bk_lenght = float(input("Введите длину одного блок-контейнера, м >> "))
bk_width = float(input("Введите ширину одного блок-контейнера, м >> "))
bk_height = float(input("Введите высоту одного блок-контейнера, м >> "))
bk_thickness_wall = float(input("Введите толщину стены одного блок-контейнера, м >> "))
bk_thickness_floor = float(input("Введите толщину пола одного блок-контейнера, м >> "))
bk_thickness_ceiling = float(input("Введите толщину потолка одного блок-контейнера, м >> "))

#ДОПОЛНИТЕЛЬНЫЕ ПАРАМЕТРЫ

bk_wall_teploisolate = float(input("Введите толщину утеплителя стен, м >>"))

#РАСЧЕТ ПРОГРАММЫ

S_floor_single = (bk_lenght - 2 * bk_thickness_wall) * (bk_width - 2 * bk_thickness_wall)
S_ceiling_single = (bk_lenght - 2 * bk_thickness_wall) * (bk_width - 2 * bk_thickness_wall)
S_wall_single = (2 * (bk_lenght - 2 * bk_thickness_wall)) + (2 * (bk_width - 2 * bk_thickness_wall)) * (bk_height - (bk_thickness_floor + bk_thickness_ceiling))

S_floor_all = bk_count * S_floor_single
S_ceiling_all = bk_count * S_ceiling_single
S_wall_all = bk_count * S_wall_single

#ОКРУГЛЕНИЕ ВЕЛИЧИН
S_floor_single = round(S_floor_single, 2)
S_ceiling_single = round(S_ceiling_single, 2)
S_wall_single = round(S_wall_single, 2)

S_floor_all = round(S_floor_all, 2)
S_ceiling_all = round(S_ceiling_all, 2)
S_wall_all = round(S_wall_all, 2)

#ВЫВОД ПРОГРАММЫ В КОНСОЛЬ
print("\n")
print("ОСНОВНЫЕ ДАННЫЕ КОММЕРЧЕСКОГО ПРЕДЛОЖЕНИЯ")
print("Коммерческое предложение №" + commercial_number + " от " + commercial_date)
print("Название коммерческого предложения: " + commercial_name)
print("\n")
print("ОСНОВНЫЕ ПАРАМЕТРЫ БЛОК-КОНТЕЙНЕРОВ")
print("Количество блок-контейнеров: " + str(bk_count) + " шт.")
print("Длина одного блок-контейнера: " + str(bk_lenght) + " м")
print("Ширина одного блок-контейнера: " + str(bk_width) + " м")
print("Высота одного блок-контейнера: " + str(bk_height) + " м")
print("\n")
print("РАСЧЕТНЫЕ ЗНАЧЕНИЯ ДЛЯ ОДНОГО БЛОК-КОНТЕЙНЕРА")
print("Площадь пола одного блок-контейнера: " + str(S_floor_single) + " м2")
print("Площадь потолка одного блок-контейнера: " + str(S_ceiling_single) + " м2")
print("Площадь стен одного блок-контейнера: " + str(S_wall_single) + " м2")
print("\n")
print("РАСЧЕТНЫЕ ЗНАЧЕНИЯ ДЛЯ ВСЕХ БЛОК-КОНТЕЙНЕРОВ")
print("Площадь пола всех блок-контейнеров: " + str(S_floor_all) + " м2")
print("Площадь потолка всех блок-контейнеров: " + str(S_ceiling_all) + " м2")
print("Площадь стен всех блок-контейнеров: " + str(S_wall_all) + " м2")

document = Document()

document.add_heading('Коммерческое предложение №' + commercial_number + " от " + commercial_date, level=1)

records = (
    ('ФИО', 'Менеджер'),
    ('Телефон, What`s App', '8-900-500-20-30'),
    ('E-mail', 'modulnie-reshenia')
)

table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ваш персональный менеджер'
hdr_cells[1].text = ''
for qty, id in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id




document.save('Коммерческое №' + commercial_number + ' от ' + commercial_date + '.docx')

input("Нажмите любую клавишу...")